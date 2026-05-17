from docx import Document
import sys

def main():
    md_file = "Documentos Artigo/Artigo_Completo_Citado.md"
    docx_file = "Documentos Artigo/Artigo_Completo_Citado.docx"
    
    doc = Document()
    
    with open(md_file, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Check if it's a heading
        if line.isupper() and len(line) < 100:
            doc.add_heading(line, level=1)
        elif line[0].isdigit() and line[1] in (" ", "."):
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)
            
    doc.save(docx_file)
    print(f"Saved {docx_file}")

if __name__ == "__main__":
    main()
